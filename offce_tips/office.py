import openpyxl as op
import docx

# CreateNewBook 
# exelfile = op.Workbook()

# Overwrite/NewSave
# exelfile.save(fname)

# OpenBook
wb = op.load_workbook("testdata.xlsx")

# GetSheet
ws = wb["Sheet1"]

# GetFirstSheet
# ws = wb.worksheets[0]

# LoadCell
def GetKeyByIndex(index):
    map = {
        0 : 'id',
        1 : 'title',
        2 : 'state',
        3 : 'description',
        4 : 'image' 
    }
    return map[index]

values = []
for row in ws.iter_rows(min_row=3, min_col=2):
    tmp_dict = {}
    for i, cell in enumerate(row):
        key_name = GetKeyByIndex(i)
        tmp_dict[key_name] = cell.value
    values.append(tmp_dict)

print(values)

# CreateNewDocument
doc = docx.Document()

# AddText
#doc.add_paragraph("・テキスト追加")

# AddImage
doc.add_picture("default_user.png")

doc.save("test.docx")